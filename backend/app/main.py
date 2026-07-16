from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy.orm.attributes import flag_modified
from typing import List
import google.generativeai as genai
import io

from app.config import settings
from app.database import get_db, engine, Base
from app.models import User, Project, AISession, Report, UploadedDocument
from app.schemas import (
    UserCreate, UserLogin, UserResponse, Token, ProjectCreate, ProjectResponse,
    ProjectDetailResponse, AISessionCreate, AISessionResponse,
    SubmitAnswersRequest, SubmitAnswersResponse, ReportResponse,
    GenerateDocumentsRequest, GenerateDocumentsResponse, ChatRequest, ChatResponse,
    UploadedDocumentResponse, AnalyzeWithDocumentsRequest
)
from app.auth import (
    get_password_hash, verify_password, create_access_token,
    get_current_user
)
from app.document_generator import DocumentGenerator
from app.response_wrapper import success_response
from app.document_parser import DocumentParser
import os
import uuid

Base.metadata.create_all(bind=engine)

app = FastAPI(title="DocuMind AI API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origins_list,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure Gemini AI
genai.configure(api_key=settings.GEMINI_API_KEY)

# Test API key and get available models
GEMINI_MODEL = None
try:
    models = genai.list_models()
    model_names = [m.name for m in models if 'generateContent' in m.supported_generation_methods]
    if not model_names:
        print("WARNING: No models found with generateContent support")
        GEMINI_MODEL = "models/gemini-1.5-pro"  # fallback
    else:
        print(f"Available models: {model_names}")
        # Use the first available model (most reliable approach)
        GEMINI_MODEL = model_names[0]
        print(f"Using model: {GEMINI_MODEL}")
except Exception as e:
    print(f"WARNING: Failed to list Gemini models: {e}")
    print("Please check your GEMINI_API_KEY in .env file")
    # Fallback to a common model name
    GEMINI_MODEL = "models/gemini-1.5-pro"

# Create upload directory if it doesn't exist
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


# Authentication Endpoints
@app.post("/auth/register", status_code=status.HTTP_201_CREATED)
def register(user: UserCreate, db: Session = Depends(get_db)):
    existing_user = db.query(User).filter(User.email == user.email).first()
    if existing_user:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Email already registered"
        )
    
    db_user = User(
        email=user.email,
        full_name=user.full_name,
        password_hash=get_password_hash(user.password)
    )
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    return success_response(data={
        "id": str(db_user.id),
        "email": db_user.email,
        "full_name": db_user.full_name,
        "created_at": db_user.created_at.isoformat()
    }, message="User registered successfully")


@app.post("/auth/login")
def login(user: UserLogin, db: Session = Depends(get_db)):
    db_user = db.query(User).filter(User.email == user.email).first()
    if not db_user or not verify_password(user.password, db_user.password_hash):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid credentials"
        )
    
    access_token = create_access_token(data={"sub": db_user.email})
    return success_response(data={
        "access_token": access_token,
        "token_type": "bearer"
    }, message="Login successful")


@app.post("/auth/logout")
def logout(current_user: User = Depends(get_current_user)):
    return success_response(message="Logged out successfully")


# Project Endpoints
@app.get("/projects")
def get_projects(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    projects = db.query(Project).filter(Project.user_id == current_user.id).order_by(Project.created_at.desc()).all()
    projects_data = [{
        "id": str(p.id),
        "name": p.name,
        "industry": p.industry,
        "description": p.description,
        "created_at": p.created_at.isoformat()
    } for p in projects]
    return success_response(data={"projects": projects_data}, message="Projects retrieved successfully")


@app.post("/projects", status_code=status.HTTP_201_CREATED)
def create_project(project: ProjectCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    db_project = Project(
        user_id=current_user.id,
        name=project.name,
        industry=project.industry,
        description=project.description
    )
    db.add(db_project)
    db.commit()
    db.refresh(db_project)
    return success_response(data={
        "id": str(db_project.id),
        "name": db_project.name,
        "industry": db_project.industry,
        "description": db_project.description,
        "created_at": db_project.created_at.isoformat()
    }, message="Project created successfully")


@app.get("/projects/{project_id}")
def get_project(project_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    return success_response(data={
        "id": str(project.id),
        "name": project.name,
        "industry": project.industry,
        "description": project.description,
        "created_at": project.created_at.isoformat()
    }, message="Project retrieved successfully")


@app.delete("/projects/{project_id}")
def delete_project(project_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    
    db.delete(project)
    db.commit()
    return success_response(message="Project deleted successfully")


@app.get("/projects/{project_id}/session")
def get_project_session(
    project_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    
    session = db.query(AISession).filter(
        AISession.project_id == project_id
    ).order_by(AISession.created_at.desc()).first()
    
    if not session:
        return success_response(data=None, message="No session found")
        
    # Extract questions or analysis from the latest assistant message
    assistant_messages = [msg for msg in session.messages if msg.get("role") == "assistant"]
    
    questions = []
    analysis = None
    
    if assistant_messages:
        last_assistant_msg = assistant_messages[-1].get("content", "")
        import json
        import re
        
        if session.status == "completed":
            try:
                json_match = re.search(r'\{.*\}', last_assistant_msg, re.DOTALL)
                if json_match:
                    analysis = json.loads(json_match.group())
                else:
                    analysis = json.loads(last_assistant_msg)
            except Exception:
                analysis = {"raw": last_assistant_msg}
        else:
            try:
                json_match = re.search(r'\{.*\}', last_assistant_msg, re.DOTALL)
                if json_match:
                    parsed = json.loads(json_match.group())
                    questions = parsed.get("questions", [])
                else:
                    parsed = json.loads(last_assistant_msg)
                    questions = parsed.get("questions", [])
            except Exception:
                questions = [q.strip() for q in last_assistant_msg.split('\n') if q.strip() and q[0].isdigit()]
                
    return success_response(data={
        "session_id": str(session.id),
        "status": session.status,
        "business_description": project.description or "",
        "questions": questions,
        "analysis": analysis
    }, message="Session retrieved successfully")


def extract_text_from_file(file: UploadFile) -> str:

    filename = file.filename.lower()
    content = ""
    if filename.endswith(('.txt', '.md', '.csv', '.json')):
        try:
            content = file.file.read().decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"Error reading text file: {e}")
    elif filename.endswith('.docx'):
        try:
            import io
            from docx import Document
            file_bytes = file.file.read()
            doc_file = io.BytesIO(file_bytes)
            doc = Document(doc_file)
            paragraphs_text = [p.text for p in doc.paragraphs if p.text]
            content = "\n".join(paragraphs_text)
        except Exception as e:
            print(f"Error reading docx file: {e}")
    elif filename.endswith('.pdf'):
        try:
            import io
            import PyPDF2
            file_bytes = file.file.read()
            pdf_file = io.BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(pdf_file)
            content = ""
            for page in reader.pages:
                content += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error reading pdf file: {e}")
    return content


# Document Upload Endpoints
@app.post("/projects/{project_id}/upload")
def upload_document(
    project_id: str,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    
    # Validate file type
    allowed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.md', '.csv', '.json'}
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type {file_ext} not allowed. Allowed types: {', '.join(allowed_extensions)}"
        )
    
    # Save file
    file_id = str(uuid.uuid4())
    file_path = os.path.join(UPLOAD_DIR, f"{file_id}_{file.filename}")
    
    try:
        with open(file_path, "wb") as buffer:
            content = file.file.read()
            buffer.write(content)
        
        # Extract text from file
        file.file.seek(0)
        extracted_text = extract_text_from_file(file)
        
        # Save to database
        uploaded_doc = UploadedDocument(
            project_id=project_id,
            filename=file.filename,
            file_type=file_ext[1:],  # Remove the dot
            file_size=len(content),
            file_path=file_path,
            extracted_text=extracted_text
        )
        db.add(uploaded_doc)
        db.commit()
        db.refresh(uploaded_doc)
        
        return success_response(
            data=UploadedDocumentResponse.model_validate(uploaded_doc),
            message="Document uploaded successfully"
        )
    except Exception as e:
        # Clean up file if database save fails
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload document: {str(e)}"
        )


@app.get("/projects/{project_id}/documents")
def list_documents(
    project_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    
    documents = db.query(UploadedDocument).filter(
        UploadedDocument.project_id == project_id
    ).order_by(UploadedDocument.uploaded_at.desc()).all()
    
    return success_response(
        data={
            "documents": [UploadedDocumentResponse.model_validate(doc) for doc in documents]
        },
        message="Documents retrieved successfully"
    )


@app.delete("/projects/{project_id}/documents/{document_id}")
def delete_document(
    project_id: str,
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    
    document = db.query(UploadedDocument).filter(
        UploadedDocument.id == document_id,
        UploadedDocument.project_id == project_id
    ).first()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    
    # Delete file from disk
    try:
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
    except Exception as e:
        print(f"Error deleting file: {e}")
    
    # Delete from database
    db.delete(document)
    db.commit()
    
    return success_response(message="Document deleted successfully")


# AI Analysis Endpoints
@app.post("/projects/{project_id}/analyze")
def start_analysis(
    project_id: str,
    business_description: str = Form(...),
    files: List[UploadFile] = File(default=[]),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    
    try:
        # Update project description in database
        project.description = business_description
        db.commit()
        
        # Extract content from uploaded files
        extracted_contexts = []
        for file in files:
            file_text = extract_text_from_file(file)
            if file_text.strip():
                extracted_contexts.append(f"--- Document: {file.filename} ---\n{file_text}")
        
        full_context = business_description
        if extracted_contexts:
            full_context += "\n\n### Attached Reference Documents:\n" + "\n\n".join(extracted_contexts)

        model = genai.GenerativeModel(GEMINI_MODEL)
        prompt = f"""
        Based on this business description and reference documents, generate 3-5 intelligent follow-up questions to gather complete context:
        
        Business Context:
        {full_context}
        
        Industry: {project.industry or 'Not specified'}
        
        Return only the questions as a numbered list.
        """
        response = model.generate_content(prompt)
        questions = [q.strip() for q in response.text.split('\n') if q.strip() and q[0].isdigit()]
        
        ai_session = AISession(
            project_id=project_id,
            messages=[
                {"role": "user", "content": full_context},
                {"role": "assistant", "content": response.text}
            ],
            status="in_progress"
        )
        db.add(ai_session)
        db.commit()
        db.refresh(ai_session)
        
        return success_response(data={
            "session_id": str(ai_session.id),
            "questions": questions,
            "messages": ai_session.messages
        }, message="Analysis started successfully")
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"AI service unavailable. Error: {str(e)}"
        )


@app.post("/projects/{project_id}/questions")
def submit_answers(
    project_id: str,
    request: SubmitAnswersRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    session = db.query(AISession).filter(
        AISession.id == request.session_id,
        AISession.project_id == project_id
    ).first()
    if not session:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Session not found")
    
    # Add answers to messages
    for qa in request.answers:
        session.messages.append({"role": "user", "content": qa.answer})
    flag_modified(session, "messages")
    
    try:
        import json
        import re
        
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        # Build conversation context
        conversation = "\n".join([f"{msg['role']}: {msg['content']}" for msg in session.messages])
        
        # Limit Q&A to a maximum of 2 rounds of assistant questions to prevent infinite loops
        assistant_messages = [m for m in session.messages if m["role"] == "assistant"]
        if len(assistant_messages) >= 2:
            session.status = "completed"
            analysis_prompt = f"""
            Generate a comprehensive business analysis based on this conversation:
            {conversation}
            
            Return as JSON with keys: workflow, pain_points, automation_opportunities, recommendations.
            """
            analysis_response = model.generate_content(analysis_prompt)
            session.messages.append({"role": "assistant", "content": analysis_response.text})
            flag_modified(session, "messages")
            
            analysis_data = {}
            try:
                json_match = re.search(r'\{.*\}', analysis_response.text, re.DOTALL)
                if json_match:
                    analysis_data = json.loads(json_match.group())
                else:
                    analysis_data = json.loads(analysis_response.text)
            except Exception:
                analysis_data = {"raw": analysis_response.text}
                
            db.commit()
            
            return success_response(data={
                "status": "completed",
                "analysis": analysis_data
            }, message="Analysis completed successfully")
        
        prompt = f"""
        Based on this conversation, determine if we have enough information to generate a comprehensive business analysis.
        If yes, generate a structured business analysis with:
        - Current workflow summary
        - Pain points
        - Automation opportunities
        - Recommended solutions
        
        If no, generate 3-5 more follow-up questions.
        
        Conversation:
        {conversation}
        
        Return your response in JSON format with keys: "status" (either "more_questions" or "completed"), 
        "questions" (array of strings if more questions needed), or "analysis" (object with the analysis if completed).
        """
        
        response = model.generate_content(prompt)
        session.messages.append({"role": "assistant", "content": response.text})
        flag_modified(session, "messages")
        
        # Try to parse JSON from Gemini response
        try:
            # Extract JSON from response (handle markdown code blocks)
            json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
            if json_match:
                parsed_response = json.loads(json_match.group())
            else:
                # Fallback: try parsing entire response
                parsed_response = json.loads(response.text)
            
            # Check if analysis is complete
            if parsed_response.get("status") == "completed":
                session.status = "completed"
                analysis = parsed_response.get("analysis", {})
                db.commit()
                
                return success_response(data={
                    "status": "completed",
                    "analysis": analysis
                }, message="Analysis completed successfully")
            elif parsed_response.get("status") == "more_questions":
                questions = parsed_response.get("questions", [])
                db.commit()
                
                return success_response(data={
                    "status": "in_progress",
                    "next_questions": questions
                }, message="More questions needed")
            else:
                # Fallback to hardcoded logic if JSON parsing fails
                if len(session.messages) > 6:
                    session.status = "completed"
                    analysis_prompt = f"""
                    Generate a comprehensive business analysis based on this conversation:
                    {conversation}
                    
                    Return as JSON with keys: workflow, pain_points, automation_opportunities, recommendations.
                    """
                    analysis_response = model.generate_content(analysis_prompt)
                    db.commit()
                    
                    return success_response(data={
                        "status": "completed",
                        "analysis": {"raw": analysis_response.text}
                    }, message="Analysis completed successfully")
                else:
                    questions = [line.strip() for line in response.text.split('\n') if line.strip() and line[0].isdigit()]
                    db.commit()
                    
                    return success_response(data={
                        "status": "in_progress",
                        "next_questions": questions
                    }, message="Answers submitted successfully")
        except json.JSONDecodeError:
            # JSON parsing failed, use fallback logic
            if len(session.messages) > 6:
                session.status = "completed"
                analysis_prompt = f"""
                Generate a comprehensive business analysis based on this conversation:
                {conversation}
                
                Return as JSON with keys: workflow, pain_points, automation_opportunities, recommendations.
                """
                analysis_response = model.generate_content(analysis_prompt)
                db.commit()
                
                return success_response(data={
                    "status": "completed",
                    "analysis": {"raw": analysis_response.text}
                }, message="Analysis completed successfully")
            else:
                questions = [line.strip() for line in response.text.split('\n') if line.strip() and line[0].isdigit()]
                db.commit()
                
                return success_response(data={
                    "status": "in_progress",
                    "next_questions": questions
                }, message="Answers submitted successfully")
            
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="AI service unavailable. Please try again later."
        )


# Document Generation Endpoints
@app.get("/projects/{project_id}/reports")
def get_reports(project_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    
    reports = db.query(Report).filter(Report.project_id == project_id).order_by(Report.created_at.desc()).all()
    reports_data = [{
        "id": str(r.id),
        "report_type": r.report_type,
        "created_at": r.created_at.isoformat(),
        "content": r.content
    } for r in reports]
    return success_response(data={"reports": reports_data}, message="Reports retrieved successfully")


@app.post("/projects/{project_id}/generate")
def generate_documents(
    project_id: str,
    request: GenerateDocumentsRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    
    if not request.document_types:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Please select at least one document type"
        )
    
    generated_reports = []
    
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        for doc_type in request.document_types:
            # Get the latest completed session
            session = db.query(AISession).filter(
                AISession.project_id == project_id,
                AISession.status == "completed"
            ).order_by(AISession.created_at.desc()).first()
            
            if not session:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="No completed AI analysis found. Please complete the analysis first."
                )
            
            conversation = "\n".join([f"{msg['role']}: {msg['content']}" for msg in session.messages])
            
            prompt = f"""
            Generate a comprehensive and professional {doc_type} based on this business analysis conversation:
            
            {conversation}
            
            Return as a professionally written document formatted in clean, plain Markdown.
            Do NOT return JSON format.
            Do NOT wrap the output in markdown code blocks (such as ```markdown or ```).
            Use clear headings (e.g., #, ##, ###), lists, bullet points, and text-based diagrams or tables where helpful.
            """
            
            response = model.generate_content(prompt)
            
            report = Report(
                project_id=project_id,
                report_type=doc_type,
                content={"raw": response.text, "structured": response.text}
            )
            db.add(report)
            db.commit()
            db.refresh(report)
            generated_reports.append({
                "id": str(report.id),
                "report_type": report.report_type,
                "created_at": report.created_at.isoformat(),
                "content": report.content
            })
        
        return success_response(data={"reports": generated_reports}, message="Documents generated successfully")
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Document generation failed. Please try again."
        )


@app.get("/reports/{report_id}/download")
def download_report(report_id: str, format: str = "pdf", current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    report = db.query(Report).join(Project).filter(
        Report.id == report_id,
        Project.user_id == current_user.id
    ).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    try:
        # Add report_type to content for document generator
        content_with_type = {
            "report_type": report.report_type,
            "raw": report.content.get("raw", "")
        }
        
        # Generate actual binary file
        file_bytes = DocumentGenerator.generate(content_with_type, format)
        
        # Determine media type
        media_types = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        }
        
        media_type = media_types.get(format.lower(), "application/octet-stream")
        
        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename={report.report_type}_{report_id}.{format}"
            }
        )
    except ValueError as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to generate document. Please try again."
        )


@app.post("/reports/{report_id}/chat")
def chat_about_report(
    report_id: str,
    request: ChatRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    report = db.query(Report).join(Project).filter(
        Report.id == report_id,
        Project.user_id == current_user.id
    ).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    try:
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        context = f"Report Content: {report.content.get('raw', '')}\n\n"
        if request.history:
            context += "Previous conversation:\n"
            for msg in request.history:
                context += f"{msg.role}: {msg.content}\n"
        
        prompt = f"{context}\nUser question: {request.message}\n\nProvide a helpful response based on the report context."
        
        response = model.generate_content(prompt)
        
        return success_response(data={"response": response.text}, message="Chat response generated successfully")
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="AI chat service unavailable. Please try again later."
        )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
